import streamlit as st
from docx import Document
import gspread
from google.oauth2.service_account import Credentials

# Google Sheets Setup
SHEET_NAME = "nogra-chemlab"  # Change this to your Google Sheet name

# Load Google Sheets credentials from Streamlit Secrets
scope = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
credentials_dict = st.secrets["google_service_account"]
creds = Credentials.from_service_account_info(dict(credentials_dict), scopes=scope)
client = gspread.authorize(creds)

# Open the Google Sheet
try:
    sheet = client.open(SHEET_NAME).sheet1
except Exception as e:
    st.error(f"Error accessing Google Sheet: {e}")
    sheet = None

# Streamlit UI
st.title("Lab Results Entry System")

# Input Fields
lab_number = st.text_input("Lab Number", key="lab_number")
sample_number = st.text_input("Sample Number", key="sample_number")
sample_description = st.text_area("Sample Description", key="sample_description")

st.subheader("Summary of Results")

# Predefined Parameters List
predefined_parameters = [
    "Moisture", "Ash", "Crude Protein", "Total Fat", "Sugar",
    "Sodium", "Potassium", "Calcium", "pH", "Water Activity",
    "Total Titrable Activity", "Carbohydrates", "Food Energy Value"
]

# Let users select multiple parameters
selected_parameters = st.multiselect("Select Parameters", predefined_parameters, key="selected_parameters")

parameters = []
for param in selected_parameters:
    parameters.append(
        {
            "Parameter": param,
            "Date Started": st.date_input(f"Date Started ({param})", key=f"date_{param}").strftime("%Y-%m-%d"),
            "Environmental Conditions": st.text_input(f"Environmental Conditions ({param})", key=f"env_{param}"),
            "Method Used": st.text_input(f"Method Used ({param})", key=f"method_{param}"),
            "Results": st.text_input(f"Results ({param})", key=f"results_{param}"),
            "Uncertainty": st.text_input(f"Uncertainty ({param})", key=f"uncertainty_{param}"),
            "Unit": st.text_input(f"Unit ({param})", key=f"unit_{param}"),
        }
    )

# Function to update Google Sheet
def update_google_sheet():
    if not sheet:
        return "Google Sheet access failed. Data not saved."
    
    data = []
    for param in parameters:
        data.append([
            lab_number, sample_number, sample_description,
            param["Parameter"], param["Date Started"], param["Environmental Conditions"],
            param["Method Used"], param["Results"], param["Uncertainty"], param["Unit"]
        ])
    
    try:
        sheet.append_rows(data)
        return "Data saved to Google Sheet!"
    except Exception as e:
        return f"Error saving to Google Sheet: {e}"

# Function to generate the report
def generate_report():
    doc = Document()
    doc.add_heading("Lab Results Report", level=1)

    doc.add_paragraph(f"Lab Number: {lab_number}")
    doc.add_paragraph(f"Sample Number: {sample_number}")
    doc.add_paragraph(f"Sample Description: {sample_description}")
    doc.add_paragraph("\nSummary of Results:")

    table = doc.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    headers = ["Parameter", "Date Started", "Environmental Conditions", "Method Used", "Results", "Uncertainty", "Unit"]

    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    for param in parameters:
        row_cells = table.add_row().cells
        row_cells[0].text = param["Parameter"]
        row_cells[1].text = param["Date Started"]
        row_cells[2].text = param["Environmental Conditions"]
        row_cells[3].text = param["Method Used"]
        row_cells[4].text = param["Results"]
        row_cells[5].text = param["Uncertainty"]
        row_cells[6].text = param["Unit"]

    report_filename = "Lab_Report.docx"
    doc.save(report_filename)

    return report_filename

# Validate required fields
def validate_fields():
    if not lab_number or not sample_number or not sample_description:
        return False, "Lab Number, Sample Number, and Sample Description are required."
    
    if not selected_parameters:
        return False, "At least one parameter must be selected."
    
    for param in parameters:
        if not param["Date Started"] or not param["Method Used"] or not param["Results"] or not param["Unit"]:
            return False, f"All fields must be filled for parameter {param['Parameter']}."
    
    return True, ""

# Save to Google Sheet & Generate Report Button
if st.button("Save to Tracker & Generate Report"):
    is_valid, error_message = validate_fields()
    
    if is_valid:
        google_sheet_status = update_google_sheet()
        file_path = generate_report()
        st.success(google_sheet_status)
        st.success("Report generated successfully!")

        with open(file_path, "rb") as file:
            st.download_button(
                "Download Report",
                file,
                file_path,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
    else:
        st.error(error_message)
